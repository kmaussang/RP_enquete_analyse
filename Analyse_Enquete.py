"""
Analyse des réponses
version : 1.2
Python 3.10

Auteur : K. Maussang
"""
from datetime import datetime
import os
import win32com.client
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import ExtractData
import Name_cleaning

current_dir=os.getcwd()
dir_fig='./Figures/'
dir_excel='./Exports-Excel/'
###########################
#   Paramètres généraux   #
###########################
IsPublic=False # rapport public ou interne
doExportpdf=True # export en .pdf
StatutQuest=8 # 8=questionnaire validé, 7=dernière page.
doAgglo1 = True # agglomère les établissements mentionnés une seule fois
###########################


# import des données
data_fn=r'./results-survey398164.xlsx'
df = pd.read_excel(data_fn)
# on supprime les colonnes inutiles
cols = [c for c in df.columns if (c[:10] != 'Durée pour' and c[:11] !='Temps total')]
df=df[cols]
# stats de remplissage
Nbm1=df[df['Dernière page']==-1].shape[0]
Nbnan=df[df['Dernière page'].isnull()].shape[0]
Nb0=df[df['Dernière page']==0].shape[0]
Nb1=df[df['Dernière page']==1].shape[0]
Nb2=df[df['Dernière page']==2].shape[0]
Nb3=df[df['Dernière page']==3].shape[0]
Nb4=df[df['Dernière page']==4].shape[0]
Nb5=df[df['Dernière page']==5].shape[0]
Nb6=df[df['Dernière page']==6].shape[0]
Nb7=df[df['Dernière page']==7].shape[0]
NbTot=df.shape[0]
NbComp=df.dropna(subset=['Date de soumission'],inplace=False).shape[0]

# data cleaning : enlever nb pages 0 et nan
df.dropna(subset=['Dernière page'], inplace=True)
df.drop(df.loc[df['Dernière page'] == 0].index, inplace=True)
# selection des donnees suivant le degres de remplissage du questionnaire
if StatutQuest==8:
    # on ne garde que les réponses completes
    df.dropna(subset=['Date de soumission'],inplace=True)
    Nb_str='Analyse sur les réponses complètes.'
else:
    df.drop(df.loc[df['Dernière page'] < StatutQuest].index, inplace=True)
    Nb_str = 'Analyse sur les réponses qui ont été au moins jusqu\'à la page '+str(StatutQuest)

# utilisation du login d'OS comme nom d'auteur
report_author=os.getlogin()

# création d'un objet 'rapport' de type document
rapport = Document()

# formatage du style 'normal' du document
style = rapport.styles['Normal']
font = style.font
font.name = 'Calisto MT'
font.size = Pt(11)

# définition des en-têtes
section = rapport.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)
header = section.header
paragraph_h = header.paragraphs[0]
if IsPublic:
    paragraph_h.text = "Enquête \"Recherches participatives\"\t\tRapport public"
else:
    paragraph_h.text = "Enquête \"Recherches participatives\"\t\tRapport interne - NE PAS DIFFUSER"
paragraph_h.style = rapport.styles["Header"]
footer = section.footer
paragraph_f = footer.paragraphs[0]
if not(IsPublic):
    paragraph_f.text = "Extraction du " + str(datetime.today())[:19] + "\t\tNE PAS DIFFUSER"
paragraph_f.style = rapport.styles["Header"]

# Titre
rapport.add_heading("Enquête \"Recherches participatives\"", level=0)
# Section 1 - paramètres d'extraction
if not(IsPublic):
    rapport.add_heading("Paramètres d'extraction", level=1)
    paragraph = rapport.add_paragraph("Rapport interne - NE PAS DIFFUSER\nExtraction réalisée par : "+report_author+'\n'+
                                        "Extraction réalisée le : "+str(datetime.today())[:10]+" à "+str(datetime.today())[11:19]+"\n"+
                                        "Source des données : "+data_fn
                                      )
    rapport.add_paragraph("Nombre d'engagements : "+str(NbTot)+'\n'+
                  "Nombre de réponses complètes : " + str(NbComp) + '\n' +
                  "Nombre de réponses jusqu'à la page -1 : " + str(Nbm1) + '\n' +
                  "Nombre de réponses sans page affichée : " + str(Nbnan) + '\n' +
                  "Nombre de réponses jusqu'à la page 0 : " + str(Nb0) + '\n' +
                  "Nombre de réponses jusqu'à la page 1 : " + str(Nb1) + '\n' +
                  "Nombre de réponses jusqu'à la page 2 : " + str(Nb2) + '\n' +
                  "Nombre de réponses jusqu'à la page 3 : " + str(Nb3) + '\n' +
                  "Nombre de réponses jusqu'à la page 4 : " + str(Nb4) + '\n' +
                  "Nombre de réponses jusqu'à la page 5 : " + str(Nb5) + '\n' +
                  "Nombre de réponses jusqu'à la page 6 : " + str(Nb6) + '\n' +
                  "Nombre de réponses jusqu'à la page 7 : " + str(Nb7)+ '\n' +
                  "Total pour vérification : " + str(Nbm1++Nbnan+Nb0+Nb1+Nb2+Nb3+Nb4+Nb5+Nb6+Nb7))

t_start=[datetime.strptime(x, "%Y-%m-%d %H:%M:%S") for x in list(df["Date de lancement"])]
t_end=[datetime.strptime(x, "%Y-%m-%d %H:%M:%S") for x in list(df["Date de la dernière action"])]
duree=[t_end[i]-t_start[i] for i in range(len(t_start))]
duree_s=[x.total_seconds() for x in duree]

fig1 = plt.figure(figsize=(12, 5), dpi=200)
plt.rcParams.update({'font.size': 11})
sns.boxplot(x="Durée (s)", data={'Durée (s)':duree_s}, width=.6)
sns.stripplot(x="Durée (s)", data={'Durée (s)':duree_s}, size=4, color=".3", linewidth=0)
plt.grid()
plt.xlabel('Durée de réponse (en s)')
plt.xlim([0,2000])
sns.despine(trim=True, left=True)
fig1.tight_layout()
fig1.savefig(dir_fig+'Duree-reponse_boxplot.png', dpi=200)

fig2 = plt.figure(dpi=200, figsize=(12, 5))
plt.rcParams.update({'font.size': 11})
plt.hist(duree_s, bins=round((np.max(duree_s) - np.min(duree_s))/60))
plt.xlabel('Durée de réponse (en s)')
plt.ylabel('Nombre de réponses')
plt.xlim([0,2000])
fig2.tight_layout()
fig2.savefig(dir_fig+'Duree-reponse_hist.png', dpi=200)


paragraph = rapport.add_paragraph(Nb_str+"\nNombre de réponses : "+str(len(df.index))+"\n"+
                                    "Durée moyenne de réponse : "+"{:.0f}".format(np.mean(duree_s)//60)+"min "+
                                    "{:.0f}".format(np.mean(duree_s) % 60)+"s\n"+
                                    "Durée minimale de réponse : "+"{:.0f}".format(np.min(duree_s)//60)+"min "+
                                    "{:.0f}".format(np.min(duree_s) % 60)+"s\n"+
                                    "Durée maximale de réponse : "+"{:.0f}".format(np.max(duree_s)//60)+"min "+
                                    "{:.0f}".format(np.max(duree_s) % 60)+"s\n"+
                                    "Pourcentage de réponses <=5min : "+"{:.0f}".format(100*len([x for x in duree_s if x<=300])/len(duree_s))+"%\n"+
                                    "Pourcentage de réponses <=10min : "+"{:.0f}".format(100*len([x for x in duree_s if x<=600])/len(duree_s))+"%\n"+
                                    "Pourcentage de réponses <=15min : "+"{:.0f}".format(100*len([x for x in duree_s if x<=900])/len(duree_s))+"%"
                                        )
paragraph.paragraph_format.space_before = Inches(0)
paragraph.paragraph_format.space_after = Inches(0)

rapport.add_picture(dir_fig+'Duree-reponse_boxplot.png', width=Inches(7.0))
last_paragraph = rapport.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
rapport.add_picture(dir_fig+'Duree-reponse_hist.png', width=Inches(7.0))
last_paragraph = rapport.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# definition des groupes de questions
groupe1=cols[8]
groupe2=cols[9:23]
groupe3=cols[23:30]
groupe4=cols[30:50]
groupe5=cols[50:63]
groupe6=cols[63:118]
groupe7=cols[118:]

#############################
#   Questions du groupe 1   #
#############################
rapport.add_heading("Groupe 1 : Type d'établissement d'appartenance", level=1)

#############################
Qn="Quelle est la nature de l'établissement auquel vous appartenez ?"
ExtractData.stat2doc(rapport,df,Qn,dir_fig+'etablissement.png',dir_excel+'etablissement.xlsx')

#############################
#   Questions du groupe 2   #
#############################
rapport.add_heading("Groupe 2 : A propos de vous...", level=1)

#############################
Qn="Type d'établissement d'appartenance"
ExtractData.stat2doc(rapport,df,Qn,dir_fig+'etablissement2.png',dir_excel+'etablissement2.xlsx')

#############################
Qn="Type d'association"
ExtractData.stat2doc(rapport,df,Qn,dir_fig+'association.png',dir_excel+'association.xlsx')

#############################
Qn="Nom de l'établissement / de l'association"
rapport.add_heading(Qn, level=2)
Qcond="Acceptez-vous de divulguer le nom de votre employeur dans l'analyse des résultats, qui sera potentiellement rendue publique ? Si vous répondez \"non\", votre établissement d'appartenance sera décrit par son secteur d'activité."
paragraph = rapport.add_paragraph('Nombre de valeurs : '+str(df[Qn].notnull().count())+"\n"+
                                'Nombre de NaN: '+str(df[Qn].isna().sum())+"\n"+
                                'Nombre de répondants ne souhaitant pas rendre public le nom de leur établissement : '+str(df[df[Qcond]=='Non'][Qn].notnull().count())+"\n"+
                                'Pourcentage de répondants ne souhaitant pas rendre public le nom de l\'établissement : '+str("{:.2f}".format(100*df[df[Qcond]=='Non'][Qn].notnull().count()/df[Qn].count()))+'%'
                              )

# nettoyage - nom de l'établissement
if doAgglo1:
    etab_str=' - >1'
else:
    etab_str=''
dict_etab=Name_cleaning.get_dict_etab()
df["Nom établissement"]=df["Nom de l'établissement / de l'association"]
for x in dict_etab.keys():
    df.loc[(df["Nom établissement"]==x),"Nom établissement"]=dict_etab[x]
# colonne supplémentaire pour agglomérer les établissements cités une seule fois
df["Nom établissement - >1"]=df["Nom établissement"]
g = df.groupby("Nom établissement - >1")["Nom établissement - >1"].transform('size')
df.loc[g == 1, "Nom établissement - >1"] = 'Etablissements mentionnés une seule fois'
# liste des établissements dont les répondants ne souhaitent pas rendre le nom public
# création d'un dictionnaire d'anonymisation
dict_etab_pub, a ={}, 0
for x in set(df[df[Qcond]=='Non']["Nom établissement"].dropna().to_list()):
    a+=1
    dict_etab_pub[x]="Etablissement X"+"{:02d}".format(a)
df["Nom établissement public"]=df["Nom établissement"]
for x in dict_etab_pub.keys():
    df.loc[(df["Nom établissement public"]==x) & (df[Qcond]=='Non'),"Nom établissement public"]=dict_etab_pub[x]
# export de la correspondance dans un fichier excel
df[["ID de la réponse","Nom de l'établissement / de l'association","Nom établissement","Nom établissement public"]].to_excel(dir_excel+'etablissement_anonymisation.xlsx')

if not(IsPublic):
    ExtractData.warning_init(rapport)
    paragraph = rapport.add_paragraph('La correspondance entre les noms d\'établissement renseignés, les noms nettoyés et les noms publics '+
                                      'est exportés dans le fichier excel '+dir_excel+'etablissement_anonymisation.xlsx\n'
                                      )
    # Statistiques sur les établissements
    df["Nom établissement"].value_counts().to_excel(dir_excel+'etablissement_stats.xlsx')
    # ajout table au rapport
    paragraph = rapport.add_paragraph(
        'Statistiques établissements - non anonymisées'
        )
    ExtractData.table2doc(df,"Nom établissement",rapport)
    ExtractData.warning_fin(rapport)

# Statistiques sur les établissements
df["Nom établissement public"].value_counts().to_excel(dir_excel+'etablissement_stats_public.xlsx')
# colonne supplémentaire pour agglomérer les établissements cités une seule fois
df["Nom établissement public - >1"]=df["Nom établissement public"]
g = df.groupby("Nom établissement public - >1")["Nom établissement public - >1"].transform('size')
df.loc[g == 1, "Nom établissement public - >1"] = 'Etablissements mentionnés une seule fois'

# ajout table au rapport
paragraph = rapport.add_paragraph(
    '\nStatistiques établissements - public'
    )
ExtractData.table2doc(df,"Nom établissement public"+etab_str,rapport)

#############################
Qn='Nature des fonctions exercées'
ExtractData.stat2doc(rapport,df,Qn,dir_fig+'fonctions_exercees.png',dir_excel+'fonctions_exercees.xlsx')

#############################
Qn='Type de contrat'
ExtractData.stat2doc(rapport,df,Qn,dir_fig+'contrat.png',dir_excel+'contrat.xlsx')

#############################
Qn='Niveau d\'étude'
ExtractData.stat2doc(rapport,df,Qn,dir_fig+'niveau_etude.png',dir_excel+'niveau_etude.xlsx')

#############################
Qn='Discipline ou thématique scientifique principale'
# nettoyage des données
dict_CNU=Name_cleaning.get_dict_CNU()
dict_CNU.update(Name_cleaning.get_dict_CNU_autres())
dict_CNU_inv=Name_cleaning.get_dict_CNU_inv()
df["CNU"]=df["Discipline ou thématique scientifique principale"]
df["CNU [Autre]"]=df["Discipline ou thématique scientifique principale [Autre]"]
for x in dict_CNU.keys():
    df.loc[df["CNU"]==x,"CNU"]=dict_CNU[x]
    df.loc[df["CNU [Autre]"] == x, "CNU"] = dict_CNU[x]
    df.loc[df["CNU [Autre]"] == x, "CNU [Autre]"]=None
paragraph = rapport.add_paragraph(
        '\nDéfinition des groupes disciplinaires CNU : '
        )
ExtractData.dict2doc(rapport,dict_CNU_inv)
ExtractData.stat2doc(rapport,df,"CNU",dir_fig+'CNU.png',dir_excel+'CNU.xlsx',title_sec='Répartition disciplinaire (groupes CNU)')
ExtractData.stat2doc(rapport,df[df['Avez-vous déjà organisé un projet de recherches participatives ?']=='Oui'],"CNU",dir_fig+'CNU_RP.png',dir_excel+'CNU_RP.xlsx',title_sec='Répartition disciplinaire (groupes CNU) des porteurs de projets de recherches participatives')
if not(IsPublic):
    ExtractData.warning_init(rapport)
    p_interne = rapport.add_paragraph('Nettoyage des réponses \'Autre\' quand cela était possible : ')
    ExtractData.dict2doc(rapport, Name_cleaning.get_dict_CNU_autres(),col_size=[100,60])
    ExtractData.warning_fin(rapport)

#############################
Qn = 'Quel est le secteur d\'activité de votre entreprise ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'secteur_activite.png', dir_excel + 'secteur_activite.xlsx')

#############################
#   Questions du groupe 3   #
#############################
rapport.add_heading("Groupe 3 : Votre rapport aux recherches participatives", level=1)

#############################
Qn='Comment évaluez-vous vous votre niveau de maîtrise des projets de recherches participatives ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'maitrise_RP.png', dir_excel + 'maitrise_RP.xlsx')

#############################
Qn='Avez-vous déjà participé, en tant que membre actif (non organisateur), à un projet de recherches participatives ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'participation_RP_non-orga.png', dir_excel + 'participation_RP_non-orga.xlsx')

#############################
Qn='Songez-vous, dans un futur proche (quelques années), participer à un projet de recherches participatives comme membre actif (non organisateur) ? Si oui, dans quel délai ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'participation_RP_non-orga_futur.png', dir_excel + 'participation_RP_non-orga_futur.xlsx')

#############################
Qn='Avez-vous déjà organisé un projet de recherches participatives ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'RP_orga.png', dir_excel + 'RP_orga.xlsx')

#############################
Qn='Avez-vous déjà organisé un projet de recherches participatives ?'
Qcond='Comment évaluez-vous vous votre niveau de maîtrise des projets de recherches participatives ?'
df.groupby(Qn)[Qcond].value_counts().to_frame().to_excel(dir_excel+'tc_RP.xlsx')
paragraph = rapport.add_paragraph('')
ExtractData.tc2doc(df,Qn,Qcond,rapport,level1_names=['Aucune connaissance sur le sujet','Notions','Maîtrise','Expert(e)'])

#############################
Qn='Songez-vous, dans un futur proche (quelques années), organiser un projet de recherches participatives ? Si oui, dans quel délai ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'RP_orga_futur.png', dir_excel + 'RP_orga_futur.xlsx')

#############################
Qn='Avez-vous déjà organisé un projet de recherches participatives ?'
Qcond='Songez-vous, dans un futur proche (quelques années), organiser un projet de recherches participatives ? Si oui, dans quel délai ?'
df.groupby(Qn)[Qcond].value_counts().to_frame().to_excel(dir_excel+'tc_RP_orga.xlsx')
paragraph = rapport.add_paragraph('')
ExtractData.tc2doc(df,Qn,Qcond,rapport,level1_names=['Non','Oui, dans les semaines/mois qui viennent','Oui, dans 1 à 2 ans','Oui, dans plus de 2 ans','Autre'])

#############################
#   Questions du groupe 4   #
#############################
rapport.add_heading("Groupe 4 : Open innovation", level=1)

#############################
Qn='Exploitez-vous des données issues de projets de recherches publiques ou de recherches participatives ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'exploitation_data.png', dir_excel + 'exploitation_data.xlsx')

#############################
Qo="Si oui, comment avez-vous eu accès à ces données ?"
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'data_access.xlsx',pub_out=IsPublic )

#############################
Qn='Votre entreprise collabore-t-elle avec des chercheurs issus des universités ou des organismes de recherche publics ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'entreprise_labo-public.png', dir_excel + 'entreprise_labo-public.xlsx')

#############################
Qo='Si oui, sous quelle(s) forme(s) ?'
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'collab_entreprise_labo-public.xlsx',pub_out=IsPublic )

#############################
Qn='Avez-vous, personnellement, collaboré avec des chercheurs issus des universités ou des organismes de recherche publics dans le cadre de vos fonctions ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'collab_perso.png', dir_excel + 'collab_perso.xlsx')

#############################
Qn="Votre entreprise possède-t-elle une stratégie d'open innovation ?"
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'OI-strat.png', dir_excel + 'OI-strat.xlsx')

#############################
Qo="Si oui, quel(s) type(s) de stratégie(s) d'open innovation ?"
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'OI-strat_type.xlsx',pub_out=IsPublic )

#############################
#   Questions du groupe 5   #
#############################
rapport.add_heading("Groupe 5 : Qualité des données", level=1)

#############################
Qn='Avez-vous rédigé et/ou utilisé un plan de gestion de données (PGD) dans le cadre de(s) projet(s) de recherches participatives que vous avez mené(s) ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'PGD.png', dir_excel + 'PGD.xlsx')

#############################
Qn='Dans le(s) projet(s) de recherches participatives que vous avez mené(s), avez-vous établi une stratégie préalable pour assurer la qualité des données ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'strat_qualite_a-priori.png', dir_excel + 'strat_qualite_a-priori.xlsx')

#############################
Qn='Si oui, laquelle ?'
ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)

#############################
Qn='Avez-vous établi un protocole de contrôle qualité a posteriori ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'cont_qualite_a-priori.png', dir_excel + 'cont_qualite_a-priori.xlsx')

#############################
Qn='Si oui, lequel ?'
ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)

#############################
Qo='Quel(s) bénéfice(s) sur vos données cette démarche de recherches participatives a-t-elle permis ?'
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'Benefice-RP_donnees.xlsx',pub_out=IsPublic)

#############################
Qn="Avez-vous pris des précautions particulières afin d'assurer la qualité des métadonnées ?"
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'qualite_metadonnees.png', dir_excel + 'qualite_metadonnees.xlsx')

#############################
Qn='Si oui, lesquelles ?'
ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)

#############################
#   Questions du groupe 6   #
#############################
rapport.add_heading("Groupe 6 : Dissémination / communication", level=1)

#############################
Qn='Avez-vous une stratégie de communication des résultats de vos activités vers le grand public, ou une démarche de vulgarisation scientifique de ces derniers ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'strat_com.png', dir_excel + 'strat_com.xlsx')

#############################
Qo='Quelles sont les cibles de votre stratégie de communication ?'
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'Cible_com.xlsx',pub_out=IsPublic )

#############################
Qo='Quels canaux de communication utilisez-vous ?'
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'Canaux_com.xlsx',pub_out=IsPublic )

#############################
Qn="Quelle(s) précaution(s) prenez-vous lors de la communication publique de vos résultats ? (message d'avertissement, notices explicatives, notes légales,...)"
ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)

#############################
Qn="Avez-vous déjà contribué ou fourni des éléments techniques à un journaliste scientifique pour la rédaction d'un article ou d'un reportage ?"
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'elements_journalistes.png', dir_excel + 'elements_journalistes.xlsx')

#############################
Qn='Publiez-vous les données issues de vos activités en open data ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'open_data.png', dir_excel + 'open_data.xlsx')

#############################
Qn='Avez-vous déjà organisé un projet de recherches participatives ?'
Qcond='Publiez-vous les données issues de vos activités en open data ?'
df.groupby(Qn)[Qcond].value_counts().to_frame().to_excel(dir_excel+'tc_RP_open-data.xlsx')
rapport.add_paragraph('')
ExtractData.tc2doc(df,Qn,Qcond,rapport,level1_names=['Non','Oui, parfois','Oui, souvent','Oui, systématiquement','Je ne sais pas'])

#############################
Qn="Quelle(s) précaution(s) prenez-vous pour accompagner la diffusion publique de vos données ? (message d'avertissement, notices explicatives, notes légales, descriptif détaillé des protocoles de collecte des données,...)"
ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)

#############################
Qo="Votre établissement travaille-t-il avec des scolaires (élèves du primaire, collégiens, lycéens, enseignants du primaire ou secondaire, IPR-IA du rectorat, IGESR) dans le cas de..."
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'etab_scolaires.xlsx',pub_out=IsPublic )

#############################
Qn='Avez-vous personnellement travaillé avec des scolaires dans le cadre de vos fonctions, y compris dans le cadre de la production de ressources pédagogiques ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'indiv_scolaires.png', dir_excel + 'indiv_scolaires.xlsx')

#############################
Qo="Votre établissement communique-t-il des ressources pédagogiques aux établissements d'enseignement supérieur ? Si oui, cochez la case correspondant aux établissements avec lesquels vous travaillez."
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'etab_ressources-pedag.xlsx',pub_out=IsPublic )

#############################
Qo='Si oui, quel(s) type(s) de ressources communiquez-vous ?'
ExtractData.tick2doc(rapport, df, Qo, dir_excel + 'ressources-pedag.xlsx',pub_out=IsPublic )

#############################
Qn='Lorsque vous publiez les données issues de vos activités en open data, leurs associez-vous une licence de réutilisation (Etalab, Creative Commons,...) ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'licences.png', dir_excel + 'licences.xlsx')

#############################
# Qn='Si oui, quelle(s) type(s) de licence(s) utilisez-vous ?'
# ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)
# aucune réponse pour le moment

#############################
#   Questions du groupe 7   #
#############################
rapport.add_heading("Groupe 7 : Contact", level=1)

#############################
Qn="Ville de l'établissement (site sur lequel vous travaillez)"
rapport.add_heading(Qn, level=2)
# nettoyage - ville de l'établissement
dict_ville=Name_cleaning.get_dict_ville()
df["Ville"]=df[Qn]
for x in dict_ville.keys():
    df.loc[(df["Ville"]==x),"Ville"]=dict_ville[x]
# export de la correspondance dans un fichier excel
df[["ID de la réponse",Qn,"Ville"]].to_excel(dir_excel+'ville-correspondance.xlsx')

if not(IsPublic):
    ExtractData.warning_init(rapport)
    paragraph = rapport.add_paragraph('La correspondance entre les noms de ville indiqués et les noms de villes corrigés'+
                                      'est exportée dans le fichier excel '+dir_excel+'ville-correspondance.xlsx'
                                      )
    ExtractData.warning_fin(rapport)
# Statistiques sur les villes
ExtractData.table2doc(df,"Ville",rapport)
df["Ville"].value_counts().to_excel(dir_excel + 'villes_stats.xlsx')

#############################
'''
BUG ?
Qn='Code postal'
rapport.add_heading(Qn, level=2)
ExtractData.table2doc(df,Qn,rapport)
df[Qn].value_counts().to_excel(dir_excel + 'CP.xlsx')
'''

#############################
Qn="Accepteriez-vous d'être contacté pour un entretien ?"
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'contact.png', dir_excel + 'contact.xlsx')

#############################
Qn='Avez-vous déjà organisé un projet de recherches participatives ?'
Qcond="Accepteriez-vous d'être contacté pour un entretien ?"
df.groupby(Qn)[Qcond].value_counts().to_frame().to_excel(dir_excel+'tc_RP_contact.xlsx')
rapport.add_paragraph('')
ExtractData.tc2doc(df,Qn,Qcond,rapport,level1_names=['Non','Oui'])

#############################
# extraction de la liste de contacts
Q01="Accepteriez-vous d'être contacté pour un entretien ?"
Q02='Avez-vous déjà organisé un projet de recherches participatives ?'
Q1="Si oui, merci d'indiquer vos coordonnées [Nom]"
Q2="Si oui, merci d'indiquer vos coordonnées [Prénom]"
Q3="Nom de l'établissement / de l'association"
Q4="Si oui, merci d'indiquer vos coordonnées [e-mail]"
Q5='Acceptez-vous de divulguer le nom de votre employeur dans l\'analyse des résultats, qui sera potentiellement rendue publique ? Si vous répondez "non", votre établissement d\'appartenance sera décrit par son secteur d\'activité.'
Q6='ID de la réponse'
contact_nom=df.loc[(df[Q01] == 'Oui') & (df[Q02]=='Oui'), Q1].dropna().to_list()
contact_prenom=df.loc[(df[Q01] == 'Oui') & (df[Q02]=='Oui'), Q2].dropna().to_list()
contact_etab=df.loc[(df[Q01] == 'Oui') & (df[Q02]=='Oui'), Q3].dropna().to_list()
contact_mail=df.loc[(df[Q01] == 'Oui') & (df[Q02]=='Oui'), Q4].dropna().to_list()
contact_divulg_employeur=df.loc[(df[Q01] == 'Oui') & (df[Q02]=='Oui'), Q5].dropna().to_list()
contact_ID=df.loc[(df[Q01] == 'Oui') & (df[Q02]=='Oui'), Q6].dropna().to_list()

contact_export = pd.DataFrame(list(zip(contact_nom, contact_prenom, contact_etab, contact_mail, contact_divulg_employeur, contact_ID)), columns=['Nom','Prénom','Etablissement','E-mail','Divulgation nom emploeur','ID'])
contact_export.to_excel(dir_excel+'liste_contacts_RPok.xlsx')
if not(IsPublic):
    ExtractData.warning_init(rapport)
    rapport.add_paragraph('La liste des répondants ayant organisé un projet de recherches participatives et accepté d\'être contacté pour un entretien '+
                                      'est exportée dans le fichier excel '+dir_excel+'liste_contacts_RPok.xlsx'
                                      )
    ExtractData.warning_fin(rapport)

#############################
Qn='Nom de votre entreprise'
ExtractData.extract_list2doc(rapport,df,Qn,pub_out=IsPublic)

#############################
Qn='Acceptez-vous de divulguer le nom de votre employeur dans l\'analyse des résultats, qui sera potentiellement rendue publique ? Si vous répondez "non", votre établissement d\'appartenance sera décrit par son secteur d\'activité.'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'Divulgation_entreprise.png', dir_excel + 'Divulgation_entreprise.xlsx')

#############################
Qn='Acceptez-vous de rendre publiques mais anonymisées vos réponses à cette enquête ? (hors nom de l\'employeur, voir question précédente). Si vous répondez "non", vos réponses ne seront pas rendues publiques.'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'Divulgation_reponses.png', dir_excel + 'Divulgation_reponses.xlsx')

#############################
Qn="Souhaitez-vous recevoir les résultats de l'enquête et son analyse par e-mail ?"
Qmail='Le cas échéant, à quelle adresse e-mail ?'
ExtractData.stat2doc(rapport, df, Qn, dir_fig + 'mailing_resultats.png', dir_excel + 'mailing_resultats.xlsx')
df.loc[df[Qn] == 'Oui', Qmail].dropna().to_excel(dir_excel+'liste_e-mail_resultats.xlsx')

#############################
Qn='Avez-vous déjà organisé un projet de recherches participatives ?'
Qcond='Acceptez-vous de divulguer le nom de votre employeur dans l\'analyse des résultats, qui sera potentiellement rendue publique ? Si vous répondez "non", votre établissement d\'appartenance sera décrit par son secteur d\'activité.'
df.groupby(Qn)[Qcond].value_counts().to_frame().to_excel(dir_excel+'tc_RP_contact.xlsx')
rapport.add_paragraph('')
ExtractData.tc2doc(df,Qn,Qcond,rapport,level1_names=['Non','Oui'])


#############################
Qn='Avez-vous rédigé et/ou utilisé un plan de gestion de données (PGD) dans le cadre de(s) projet(s) de recherches participatives que vous avez mené(s) ?'
ExtractData.stat2doc(rapport,df[df[Qn]=='Oui'],"CNU",dir_fig+'CNU_PGD.png',dir_excel+'CNU_PGD.xlsx',title_sec='Répartition disciplinaire (groupes CNU) des porteurs de projets de recherches participatives ayant utilisé un PGD')


#############################
#   sauvegarde du rapport   #
#############################
rapport.save("./Rapports/rapport_"+str(datetime.today())[:10]+".docx")

# conversion en .pdf
if doExportpdf:
    wdFormatPDF = 17
    inputFile = os.path.abspath('./Rapports/rapport_'+str(datetime.today())[:10]+'.docx')
    outputFile = os.path.abspath('./Rapports/rapport_'+str(datetime.today())[:10]+'.pdf')
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(inputFile)
    doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    # ouverture du fichier .pdf
    os.chdir(current_dir+'/Rapports')
    os.system('rapport_'+str(datetime.today())[:10]+'.pdf')
