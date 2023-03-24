"""
Extraction des listes de réponses pour certaines questions à champ de texte libre
version : 1.0
Python 3.10
Auteur : K. Maussang
"""
from datetime import datetime
import os
import win32com.client
from docx import Document
from docx.shared import Inches, Mm, Pt
import pandas as pd
import ExtractData
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

current_dir=os.getcwd()
dir_extract='./Listes-reponses/'

############################
#   Paramètres d'analyse   #
############################
StatutQuest=8 # 8=questionnaire validé, 7=dernière page.

# import des données
data_fn=r'./results-survey398164.xlsx'
df = pd.read_excel(data_fn)
# on supprime les colonnes inutiles
cols = [c for c in df.columns if (c[:10] != 'Durée pour' and c[:11] !='Temps total')]
df=df[cols]
# sélection des réponses que l'on souhaite garder
# data cleaning : enlever nb pages 0 et nan
df.dropna(subset=['Dernière page'], inplace=True)
df.drop(df.loc[df['Dernière page'] == 0].index, inplace=True)
# selection des donnees suivant le degres de remplissage du questionnaire
if StatutQuest==8:
    # on ne garde que les réponses completes
    df.dropna(subset=['Date de soumission'],inplace=True)
else:
    df.drop(df.loc[df['Dernière page'] < StatutQuest].index, inplace=True)

def get_answers(Qn):
    q1='ID de la réponse'
    q2 = 'Acceptez-vous de rendre publiques mais anonymisées vos réponses à cette enquête ? (hors nom de l\'employeur, voir question précédente). Si vous répondez "non", vos réponses ne seront pas rendues publiques.'
    q3='Si oui, merci d\'indiquer vos coordonnées [Nom]'
    q4='Si oui, merci d\'indiquer vos coordonnées [Prénom]'
    q5='Si oui, merci d\'indiquer vos coordonnées [e-mail]'
    q6="Accepteriez-vous d'être contacté pour un entretien ?"
    q7="Nom de l'établissement / de l'association"
    out.add_paragraph('Nombre de réponses : ' + str(df[Qn].notnull().sum()) + '\n'
                                                                              'Nombre de réponses publiques : ' + str(df[df[q2] == 'Oui'][Q2].notnull().sum()))

    liste_ID = ExtractData.extract_list(df,[Qn, q1, q2, q3, q4, q5, q6, q7])
    for i in range(len(liste_ID[1])):
        out_str = liste_ID[0][i] +'\n' +'-'*100 +'\n\t\tID : ' + str(liste_ID[1][i]) + '\n\t\tRéponse publique : ' + str(
            liste_ID[2][i]) + '\n\t\t' + str(liste_ID[4][i]) + ' ' + str(liste_ID[3][i]) + ' (' + str(
            liste_ID[5][i]) + ')\n\t\t' + str(liste_ID[7][i]) + '\n\t\tAccepte d\'être contacté(e) : ' + str(
            liste_ID[6][i]) + '\n'+'-'*100+'\n'
        paragraph = out.add_paragraph(out_str, style='List Bullet')
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)

out = Document()
# formatage du style 'normal' du document
style = out.styles['Normal']
font = style.font
font.name = 'Calisto MT'
font.size = Pt(11)
# définition des en-têtes
section = out.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)
header = section.header
paragraph_h = header.paragraphs[0]
paragraph_h.text = "Enquête \"Recherches participatives\"\tSynthèse\tNE PAS DIFFUSER"
paragraph_h.style = out.styles["Header"]
footer = section.footer
paragraph_f = footer.paragraphs[0]
paragraph_f.text = "Extraction du " + str(datetime.today())[:19] + "\t\tNE PAS DIFFUSER"
paragraph_f.style = out.styles["Header"]
# Titre
out.add_heading("Synthèse des réponses pour les questions à champ de texte libre", level=0)

Q1='Dans le(s) projet(s) de recherches participatives que vous avez mené(s), avez-vous établi une stratégie préalable pour assurer la qualité des données ?'
Q2='Si oui, laquelle ?'
out.add_heading(Q1, level=1)
get_answers(Q2)

Q1='Avez-vous établi un protocole de contrôle qualité a posteriori ?'
Q2='Si oui, lequel ?'
out.add_heading(Q1, level=1)
get_answers(Q2)

Q1='Quel(s) bénéfice(s) sur vos données cette démarche de recherches participatives a-t-elle permis ?'
out.add_heading(Q1, level=1)
out_cell, nb, ID, reponse = [], [], [], []
for x in df.columns:
    if (Q1 in x and x.replace(Q1 + ' ', '')[1:-1] != 'Autre'):
        out_cell.append(x.replace(Q1 + ' ', '')[1:-1])
        nb.append(df[df[x] == 'Oui'][x].shape[0])
        ID.extend(df[df[x] == 'Oui']['ID de la réponse'].to_list())
        reponse.extend([x.replace(Q1 + ' ', '')[1:-1]])
table = out.add_table(rows=len(out_cell), cols=3)
for i in range(len(out_cell)):
    table.cell(i, 0).text = str(out_cell[i])
    table.cell(i, 1).text = str(nb[i])
    table.cell(i, 2).text = "{:.1f}%".format(100 * nb[i] / len(set(ID)))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
out.add_paragraph('\nTraitement des réponses \'Autres\'')
get_answers(Q1+' [Autre]')

Q1="Avez-vous pris des précautions particulières afin d'assurer la qualité des métadonnées ?"
Q2='Si oui, lesquelles ?'
out.add_heading(Q1, level=1)
get_answers(Q2)

Q1="Quelle(s) précaution(s) prenez-vous lors de la communication publique de vos résultats ? (message d'avertissement, notices explicatives, notes légales,...)"
out.add_heading(Q1, level=1)
get_answers(Q1)

Q1="Quelle(s) précaution(s) prenez-vous pour accompagner la diffusion publique de vos données ? (message d'avertissement, notices explicatives, notes légales, descriptif détaillé des protocoles de collecte des données,...)"
out.add_heading(Q1, level=1)
get_answers(Q1)

out.save(dir_extract + "Synthese-des-reponses-texte-libre.docx")
wdFormatPDF = 17
inputFile = os.path.abspath(dir_extract + "Synthese-des-reponses-texte-libre.docx")
outputFile = os.path.abspath(dir_extract + "Synthese-des-reponses-texte-libre.pdf")
word = win32com.client.Dispatch('Word.Application')
doc = word.Documents.Open(inputFile)
doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()
# ouverture du fichier .pdf
os.chdir(dir_extract)
os.system('Synthese-des-reponses-texte-libre.pdf')
