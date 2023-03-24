"""
Extraction des réponses par ID
version : 1.0
Python 3.10
Auteur : K. Maussang
"""
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches, Mm, Pt
import pandas as pd

current_dir=os.getcwd()
dir_extract='./Fiches-individuelles/'
dir_extract_synthese='./Fiches-individuelles-Synthese/'
###########################
# import des données
data_fn=r'./results-survey398164.xlsx'
df = pd.read_excel(data_fn)
# on supprime les colonnes inutiles
cols = [c for c in df.columns if (c[:10] != 'Durée pour' and c[:11] !='Temps total')]
df=df[cols]
# on ne garde que les réponses completes
df.dropna(subset=['Date de soumission'],inplace=True)

# definition des groupes de questions
groupe1=cols[8]
groupe2=cols[9:23]
groupe3=cols[23:30]
groupe4=cols[30:50]
groupe5=cols[50:63]
groupe6=cols[63:118]
groupe7=cols[118:]


for x in df['ID de la réponse'].to_list():
    print('Création de la fiche de synthèse n°'+str(x))
    fiche = Document()
    # formatage du style 'normal' du document
    style = fiche.styles['Normal']
    font = style.font
    font.name = 'Calisto MT'
    font.size = Pt(11)
    # définition des en-têtes
    section = fiche.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    header = section.header
    paragraph_h = header.paragraphs[0]
    paragraph_h.text = "Enquête \"Recherches participatives\"\tSynthèse individuelle\tNE PAS DIFFUSER"
    paragraph_h.style = fiche.styles["Header"]
    footer = section.footer
    paragraph_f = footer.paragraphs[0]
    paragraph_f.text = "Extraction du " + str(datetime.today())[:19] + "\t\tNE PAS DIFFUSER"
    paragraph_f.style = fiche.styles["Header"]
    # Titre
    fiche.add_heading("ID : "+str(x), level=0)
    fiche.add_paragraph('Date de soumission : '+str(df[df['ID de la réponse']==x]['Date de soumission'].to_list()[0]))
    fiche.add_paragraph('Nom : '+str(df[df['ID de la réponse'] == x]['Si oui, merci d\'indiquer vos coordonnées [Nom]'].to_list()[0])+'\n'+
                        'Prénom : '+str(df[df['ID de la réponse'] == x]['Si oui, merci d\'indiquer vos coordonnées [Prénom]'].to_list()[0])+'\n'+
                        'E-mail : '+str(df[df['ID de la réponse'] == x]['Si oui, merci d\'indiquer vos coordonnées [e-mail]'].to_list()[0])+'\n'+
                        "Accepteriez-vous d'être contacté pour un entretien ? : "+str(df[df['ID de la réponse'] == x]["Accepteriez-vous d'être contacté pour un entretien ?"].to_list()[0]))
    fiche.add_paragraph('Type d\'établissement d\'appartenance : '+str(df[df['ID de la réponse'] == x]['Type d\'établissement d\'appartenance'].to_list()[0])+'\n'
                        'Nom de l\'établissement / de l\'association : '+str(df[df['ID de la réponse'] == x]['Nom de l\'établissement / de l\'association'].to_list()[0])+'\n'
                        'Nature des fonctions exercées : '+str(df[df['ID de la réponse'] == x]['Nature des fonctions exercées'].to_list()[0]))
    fiche.add_paragraph('Ville : '+str(df[df['ID de la réponse'] == x]['Ville de l\'établissement (site sur lequel vous travaillez)'].to_list()[0])+'\n'
                        'Code postal : '+"{:.0f}".format(df[df['ID de la réponse'] == x]['Code postal'].to_list()[0]))
    fiche.add_paragraph('Discipline ou thématique scientifique principale : ' + str(df[df['ID de la réponse'] == x]['Discipline ou thématique scientifique principale'].to_list()[0]) )
    fiche.add_paragraph('Niveau de maîtrise en recherches participatives : '+str(df[df['ID de la réponse'] == x]['Comment évaluez-vous vous votre niveau de maîtrise des projets de recherches participatives ?'].to_list()[0])+'\n'
                        'Avez-vous déjà participé à un projet de recherches participatives (non organisateur) ? '+str(df[df['ID de la réponse'] == x]['Avez-vous déjà participé, en tant que membre actif (non organisateur), à un projet de recherches participatives ?'].to_list()[0])+'\n'
                        'Avez-vous déjà organisé un projet de recherches participatives ? '+str(df[df['ID de la réponse'] == x]['Avez-vous déjà organisé un projet de recherches participatives ?'].to_list()[0]))
    fiche.add_paragraph('Publiez-vous les données issues de vos activités en open data ? ' + str(df[df['ID de la réponse'] == x]['Publiez-vous les données issues de vos activités en open data ?'].to_list()[0]) )

    fiche.add_paragraph('Acceptez-vous de divulguer le nom de votre employeur dans l\'analyse des résultats ? '+str(df[df['ID de la réponse'] == x]['Acceptez-vous de divulguer le nom de votre employeur dans l\'analyse des résultats, qui sera potentiellement rendue publique ? Si vous répondez "non", votre établissement d\'appartenance sera décrit par son secteur d\'activité.'].to_list()[0])+'\n'
                        'Acceptez-vous de rendre publiques mais anonymisées vos réponses à cette enquête ? '+str(df[df['ID de la réponse'] == x]['Acceptez-vous de rendre publiques mais anonymisées vos réponses à cette enquête ? (hors nom de l\'employeur, voir question précédente). Si vous répondez "non", vos réponses ne seront pas rendues publiques.'].to_list()[0])+'\n'
                        'Souhaitez-vous recevoir les résultats de l\'enquête et son analyse par e-mail ? '+str(df[df['ID de la réponse'] == x]['Souhaitez-vous recevoir les résultats de l\'enquête et son analyse par e-mail ?'].to_list()[0]))

    fiche.add_paragraph('Avez-vous rédigé et/ou utilisé un PGD dans le cadre de(s) projet(s) de recherches participatives que vous avez mené(s) ? ' + str(df[df['ID de la réponse'] == x]['Avez-vous rédigé et/ou utilisé un plan de gestion de données (PGD) dans le cadre de(s) projet(s) de recherches participatives que vous avez mené(s) ?'].to_list()[0]) )
    fiche.add_paragraph('Dans le(s) projet(s) de recherches participatives que vous avez mené(s), avez-vous établi une stratégie préalable pour assurer la qualité des données ? ' + str(df[df['ID de la réponse'] == x]['Dans le(s) projet(s) de recherches participatives que vous avez mené(s), avez-vous établi une stratégie préalable pour assurer la qualité des données ?'].to_list()[0])
                        +'\n'+str(df[df['ID de la réponse'] == x]['Si oui, laquelle ?'].to_list()[0]))
    fiche.add_paragraph('Avez-vous établi un protocole de contrôle qualité a posteriori ? ' + str(df[df['ID de la réponse'] == x]['Avez-vous établi un protocole de contrôle qualité a posteriori ?'].to_list()[0])
                        +'\n'+str(df[df['ID de la réponse'] == x]['Si oui, lequel ?'].to_list()[0]))
    fiche.add_paragraph('Avez-vous pris des précautions particulières afin d\'assurer la qualité des métadonnées ? ' + str(df[df['ID de la réponse'] == x]['Avez-vous pris des précautions particulières afin d\'assurer la qualité des métadonnées ?'].to_list()[0])
                        +'\n'+str(df[df['ID de la réponse'] == x]['Si oui, lesquelles ?'].to_list()[0]))
    fiche.save(dir_extract_synthese + 'ID-'+str(x)+"_fiche_synthese.docx")

for x in df['ID de la réponse'].to_list():
    print('Création de la fiche n°'+str(x))
    fiche = Document()
    # formatage du style 'normal' du document
    style = fiche.styles['Normal']
    font = style.font
    font.name = 'Calisto MT'
    font.size = Pt(11)
    # définition des en-têtes
    section = fiche.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    header = section.header
    paragraph_h = header.paragraphs[0]
    paragraph_h.text = "Enquête \"Recherches participatives\"\tFiche individuelle\tNE PAS DIFFUSER"
    paragraph_h.style = fiche.styles["Header"]
    footer = section.footer
    paragraph_f = footer.paragraphs[0]
    paragraph_f.text = "Extraction du " + str(datetime.today())[:19] + "\t\tNE PAS DIFFUSER"
    paragraph_f.style = fiche.styles["Header"]
    # Titre
    fiche.add_heading("ID : "+str(x), level=0)
    fiche.add_paragraph('Date de soumission : '+str(df[df['ID de la réponse']==x]['Date de soumission'].to_list()[0]))

    fiche.add_heading("Groupe 1 : Type d'établissement d'appartenance", level=1)
    fiche.add_paragraph(str(df[df['ID de la réponse'] == x][groupe1].to_list()[0]))
    fiche.add_heading("Groupe 2 : A propos de vous...", level=1)
    for Q in groupe2:
        if str(df[df['ID de la réponse'] == x][Q].to_list()[0]) != 'nan':
            fiche.add_heading(Q, level=3)
            paragraph=fiche.add_paragraph(str(df[df['ID de la réponse'] == x][Q].to_list()[0]))
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
    fiche.add_heading("Groupe 3 : Votre rapport aux recherches participatives", level=1)
    for Q in groupe3:
        if str(df[df['ID de la réponse'] == x][Q].to_list()[0]) != 'nan':
            fiche.add_heading(Q, level=3)
            paragraph=fiche.add_paragraph(str(df[df['ID de la réponse'] == x][Q].to_list()[0]))
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
    fiche.add_heading("Groupe 4 : Open innovation", level=1)
    for Q in groupe4:
        if str(df[df['ID de la réponse'] == x][Q].to_list()[0]) != 'nan':
            fiche.add_heading(Q, level=3)
            paragraph=fiche.add_paragraph(str(df[df['ID de la réponse'] == x][Q].to_list()[0]))
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
    fiche.add_heading("Groupe 5 : Qualité des données", level=1)
    for Q in groupe5:
        if str(df[df['ID de la réponse'] == x][Q].to_list()[0]) != 'nan':
            fiche.add_heading(Q, level=3)
            paragraph=fiche.add_paragraph(str(df[df['ID de la réponse'] == x][Q].to_list()[0]))
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
    fiche.add_heading("Groupe 6 : Dissémination / communication", level=1)
    for Q in groupe6:
        if str(df[df['ID de la réponse'] == x][Q].to_list()[0]) != 'nan':
            fiche.add_heading(Q, level=3)
            paragraph=fiche.add_paragraph(str(df[df['ID de la réponse'] == x][Q].to_list()[0]))
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
    fiche.add_heading("Groupe 7 : Contact", level=1)
    for Q in groupe7:
        if str(df[df['ID de la réponse'] == x][Q].to_list()[0]) != 'nan':
            fiche.add_heading(Q, level=3)
            paragraph=fiche.add_paragraph(str(df[df['ID de la réponse'] == x][Q].to_list()[0]))
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)

    fiche.save(dir_extract + 'ID-'+str(x)+"_fiche.docx")
