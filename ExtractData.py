"""
Définition de fonctions utiles pour l'analyse des réponses
version : 1.0
Python 3.10
Auteur : K. Maussang
"""
import matplotlib.pyplot as plt
import pandas as pd
from docx.shared import Inches, Mm
from docx.enum.text import WD_COLOR
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

def warning_init(doc):
    p_interne = doc.add_paragraph('\n')
    run = p_interne.add_run('-'*40+' Information interne '+'-'*40)
    run.bold = True
    run.font.highlight_color = WD_COLOR.YELLOW
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def warning_fin(doc):
    p_interne = doc.add_paragraph('\n')
    run = p_interne.add_run('-'*40+' Fin - information interne '+'-'*40)
    run.bold = True
    run.font.highlight_color = WD_COLOR.YELLOW
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def pie_plot(df,Qn, start_angle=90, fig_size=(12, 5), font_size=11, label_size=12, doSave=False, fn_save='no_fn', doClose=False):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    # fonction pour représentation en pie chart d'une information
    nb_autres = len(df[df[Qn] == 'Autre'])
    # plt.subplots(figsize=fig_size)
    fig = plt.figure(figsize=fig_size, dpi=150)
    plt.rcParams.update({'font.size': font_size})
    df[Qn].value_counts().plot(kind='pie', startangle=start_angle, autopct='%.1f%%', textprops={'fontsize': label_size})
    plt.ylabel("")
    if nb_autres == 0:
        plt.xlabel('Nombre total de réponses : ' + str(df[Qn].count()))
    else:
        plt.xlabel('Nombre total de réponses : ' + str(df[Qn].count()) + ' dont ' + str(nb_autres) + " 'Autres'")
    fig.tight_layout()
    if doSave: fig.savefig(fn_save, dpi=200)
    if doClose: plt.close(fig)

def hist_plot(df, Qn, fig_size=(12, 5), font_size=11, doSave=False, fn_save='no_fn'):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    nb_autres = len(df[df[Qn] == 'Autre'])
    fig = plt.figure(figsize=fig_size, dpi=150)
    plt.rcParams.update({'font.size': font_size})
    df[Qn].value_counts().plot(kind='bar')
    if nb_autres == 0:
        plt.xlabel('Nombre total de réponses : ' + str(df[Qn].count()))
    else:
        plt.xlabel('Nombre total de réponses : ' + str(df[Qn].count()) + ' dont ' + str(nb_autres) + " 'Autres'")
    fig.tight_layout()
    if doSave: fig.savefig(fn_save, dpi=200)

def integrate_autre(df, Qn):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    x = pd.DataFrame(df[df[Qn] != 'Autre'][Qn])
    x.reset_index(drop=True, inplace=True)
    y = pd.DataFrame(df[df[Qn] == 'Autre'][[Qn + ' [Autre]']])
    y.rename(columns={Qn + ' [Autre]': Qn}, inplace=True)
    y.reset_index(drop=True, inplace=True)
    x = x.append(tmp2, ignore_index=True)
    return x

def extract_list_autres(df, Qn):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    return df[Qn + ' [Autre]'].dropna().to_list()

def extract_list(df, Qn):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    if len(Qn)==1:
        return df[Qn[0]].dropna().to_list()
    else:
        res=[]
        df_red=df.dropna(subset=[Qn[0]])
        for x in Qn:
            li = df_red[x].to_list()
            res.append(li)
        return res


def display_autre(df, Qn):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    if Qn + ' [Autre]' in df.columns:
        list_autres = extract_list_autres(Qn)
        if len(list_autres) == 0:
            print("\nPas d\'entrée 'Autre'")
            print('-' * 100)
        elif len(list_autres) == 1:
            print("\nUne seule entrée 'Autre'\n" + '-' * 50)
            print(list_autres[0])
            print('-' * 100)
        else:
            print("\nListe des entrées de type 'Autre' (" + str(len(list_autres)) + ' entrées)\n' + '-' * 50)
            for x in list_autres: print(x)
            print('-' * 100)
    else:
        print('N/A - ne s\'applique pas à cette question.')

def get_stat(df, Qn):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    print('-' * 19 + '\n' + '| Résultats bruts |\n' + '-' * 19)
    print(df[Qn].value_counts().to_frame())
    print('-' * 100)
    print('-' * 19 + '\n' + '| Répartition (%) |\n' + '-' * 19)
    print(100 * df[Qn].value_counts(normalize=True).to_frame())
    print('-' * 100)

def get_stat_cond(df, Qn, Qp, Cond):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    # Qn = intitulé de la colonne conditionnelle
    # Cond = condition
    df_tmp = df[df[Qp] == Cond]
    print('-' * 19 + '\n' + '| Résultats bruts |\n' + '-' * 19)
    print(df_tmp[Qn].value_counts().to_frame())
    print('-' * 100)
    print('-' * 19 + '\n' + '| Répartition (%) |\n' + '-' * 19)
    print(100 * df_tmp[Qn].value_counts(normalize=True).to_frame())
    print('-' * 100)

def get_stat_tick(df, Qo):
    # df = dataframe de travail
    # Qo = intitulé de la colonne de travail
    out, out_nb = [], []
    for x in df.columns:
        if (Qo in x and x.replace(Qo + ' ', '')[1:-1] != 'Autre'):
            out.append(x.replace(Qo + ' ', '')[1:-1])
            out_nb.append(df[df[x] == 'Oui'][x].shape[0])
    print('-' * 19 + '\n' + '| Résultats bruts |\n' + '-' * 19)
    for i in range(len(out)):
        print(out[i] + ' ' * (3 + len(max(out, key=len)) - len(out[i])) + str(out_nb[i]))
    print('-' * 100)
    print('-' * 19 + '\n' + '| Répartition (%) |\n' + '-' * 19)
    for i in range(len(out)):
        print(out[i] + ' ' * (3 + len(max(out, key=len)) - len(out[i])) + str(
            "{:.2f}".format(100 * out_nb[i] / sum(out_nb))) + '%')
    print('-' * 100)

def table2doc(df,Qn,doc,doPC=True):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    # Qcond = colonne conditionnelle
    # doc = nom de l'objet texte pour l'export
    # doPC = True si l'on souhaite rajouter les pourcentages
    table_out=df[Qn].value_counts().to_frame() # export sous forme d'un dataframe
    table_out.insert(0,'Index',table_out.index,True) # ajout de la colonne des index
    if doPC:
        PC = 100 * table_out[Qn] / table_out[Qn].sum()
        table_out.insert(2,'PC',PC.map("{:.2f}%".format),True)
    table = doc.add_table(rows=(table_out.shape[0]), cols=table_out.shape[1])
    for i, column in enumerate(table_out):
        for row in range(table_out.shape[0]):
            table.cell(row, i).text = str(table_out[column][row])
    table.style = 'Table Grid'
    table.allow_autofit = True
    table.autofit = True
    if (Qn + ' [Autre]' in df.columns):
        if len(extract_list_autres(df, Qn)) != 0:
            paragraph = doc.add_paragraph("\nListe des réponses \"Autres\" :")
            for x in extract_list_autres(df, Qn):
                paragraph = doc.add_paragraph(str(x), style='List Bullet')
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)


def tc2doc(df,Qn,Qcond,doc,level1_names=None):
    # df = dataframe de travail
    # Qn = intitulé de la colonne de travail
    # Qcond = colonne conditionnelle
    # doc = nom de l'objet texte pour l'export
    # level1_names = liste des intitulés du niveau 1 (réponses de la colonne conditionnelle)
    # permet de définir un ordre à la liste sinon généré automatiquement

    # tc = serie pandas avec le nombre d'occurences
    # tc.index.levels[i] = descripteurs de niveau i
    # tc.keys().names = intitulé des colonnes qui ont construit tc
    tc = df.groupby(Qn)[Qcond].value_counts()
    if level1_names==None: level1_names = tc.index.levels[1].to_list()
    table = doc.add_table(rows=2 * len(level1_names) + 1, cols=len(tc.keys().names) + 2)
    for i in range(len(tc.keys().names)):
        table.cell(0, i).text = tc.keys().names[i]
    table.cell(0, len(tc.keys().names)).text = "Nb répondants"
    table.cell(0, len(tc.keys().names) + 1).text = "%"
    a = 1
    for x in tc.index.levels[0].to_list():
        table.cell(a, 0).text = x
        b = len(level1_names)
        # fusion des cellules des descripteurs de niveau 0
        table.cell(a, 0).merge(table.cell(a + b - 1, 0))
        table.cell(a, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # somme des valeurs pour la valeur x en niveau 0
        n_tot = tc.xs(x, level=0, axis=0, drop_level=False).sum()
        c = 0
        for y in level1_names:
            table.cell(a + c, 1).text = y
            if y in tc.xs(x, level=0, axis=0, drop_level=True).keys().to_list():
                # nombre de valeurs pour la valeur x en niveau 0 et y en niveau 1
                n = tc.xs(x, level=0, axis=0, drop_level=False).xs(y, level=1, axis=0, drop_level=False)[0]
            else:
                n=0
            table.cell(a + c, 2).text = str(n)
            table.cell(a + c, 3).text = "{:.1f}%".format(100 * n / n_tot)
            c += 1
        a += b
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def pic2doc(doc,pic_fn):
    doc.add_picture(pic_fn, width=Inches(7.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def stat2doc(doc,df,Qn,fn_fig,fn_excel,title_sec=None):
    if title_sec==None: doc.add_heading(Qn, level=2)
    else: doc.add_heading(title_sec, level=2)
    pie_plot(df, Qn, doSave=True, fn_save=fn_fig, doClose=True)
    pic2doc(doc, fn_fig)
    # export excel
    df[Qn].value_counts().to_excel(fn_excel)
    # ajout table au rapport
    table2doc(df, Qn, doc)

def tick2doc(doc,df,Qo,fn_excel,pub_out=False,title_sec=None):
    if title_sec==None: doc.add_heading(Qo, level=2)
    else: doc.add_heading(title_sec, level=2)
    Qpub = 'Acceptez-vous de rendre publiques mais anonymisées vos réponses à cette enquête ? (hors nom de l\'employeur, voir question précédente). Si vous répondez "non", vos réponses ne seront pas rendues publiques.'
    out, out_nb, ID, reponse = [], [], [], []
    for x in df.columns:
        if (Qo in x and x.replace(Qo + ' ', '')[1:-1] != 'Autre'):
            out.append(x.replace(Qo + ' ', '')[1:-1])
            out_nb.append(df[df[x] == 'Oui'][x].shape[0])
            ID.extend(df[df[x] == 'Oui']['ID de la réponse'].to_list())
            reponse.extend([x.replace(Qo + ' ', '')[1:-1]])
    table = doc.add_table(rows=len(out), cols=3)
    for i in range(len(out)):
        table.cell(i, 0).text = str(out[i])
        table.cell(i, 1).text = str(out_nb[i])
        table.cell(i, 2).text = "{:.1f}%".format(100 * out_nb[i]/len(set(ID)))
    if (Qo + ' [Autre]' in df.columns):
        if len(extract_list_autres(df, Qo)) != 0:
            table.cell(i, 0).text = 'Autre'
            table.cell(i, 1).text = str(len(extract_list_autres(df, Qo)))
            table.cell(i, 2).text = "{:.1f}%".format(100 * len(extract_list_autres(df, Qo)) / len(set(ID)))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    paragraph = doc.add_paragraph("Nombre total de réponses : "+str(len(set(ID))))
    if (Qo + ' [Autre]' in df.columns):
        if len(extract_list_autres(df, Qo)) != 0:
            paragraph = doc.add_paragraph("Liste des réponses \"Autres\" :")
            for x in extract_list_autres(df[df[Qpub]=='Oui'], Qo):
                paragraph = doc.add_paragraph(str(x), style='List Bullet')
            if (len(df[df[Qpub]=='Non'][Qo + ' [Autre]'].dropna().to_list()) != 0):
                paragraph = doc.add_paragraph(str(len(df[df[Qpub]=='Non'][Qo + ' [Autre]'].dropna().to_list()))+' réponses que les répondants ne souhaitent pas rendre publiques', style='List Bullet')
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
            out_nb.append(len(extract_list_autres(df, Qo)))
            reponse.append('Autre')
    df_export = pd.DataFrame(out_nb, index=reponse, columns=[Qo])
    df_export.to_excel(fn_excel)
    if not(pub_out):
        warning_init(doc)
        liste_ID = extract_list(df, [Qo + ' [Autre]', 'ID de la réponse', Qpub])
        for i in range(len(liste_ID[0])):
            out_str = liste_ID[0][i] + ' (ID : ' + str(liste_ID[1][i]) + ', réponse publique : ' + liste_ID[2][i] + ')'
            doc.add_paragraph(out_str, style='List Bullet')
        warning_fin(doc)

def dict2doc(doc,dict_in,col_size=[20,150]):
    table = doc.add_table(rows=len(dict_in), cols=2)
    row=list(dict_in.keys())
    for x in range(len(row)):
        table.cell(x, 0).text = str(row[x])
        table.cell(x, 1).text = str(dict_in[row[x]])
    table.style = 'Table Grid'
    for cell in table.columns[0].cells:
        cell.width = Mm(col_size[0])
    for cell in table.columns[1].cells:
        cell.width = Mm(col_size[1])
    table.allow_autofit = True
    table.autofit = True

def extract_list2doc(doc,df,Qn,pub_out=False,title_sec=None):
    if title_sec==None: doc.add_heading(Qn, level=2)
    else: doc.add_heading(title_sec, level=2)
    Qpub='Acceptez-vous de rendre publiques mais anonymisées vos réponses à cette enquête ? (hors nom de l\'employeur, voir question précédente). Si vous répondez "non", vos réponses ne seront pas rendues publiques.'
    for x in extract_list(df[df[Qpub]=='Oui'], [Qn]):
        doc.add_paragraph(x, style='List Bullet')
    if (len(df[df[Qpub] == 'Non'][Qn].dropna().to_list()) != 0):
        doc.add_paragraph(str(len(df[df[Qpub] == 'Non'][Qn].dropna().to_list())) + ' réponses que les répondants ne souhaitent pas rendre publiques', style='List Bullet')
    if not (pub_out):
        warning_init(doc)
        liste_ID = extract_list(df, [Qn, 'ID de la réponse', Qpub])
        for i in range(len(liste_ID[0])):
            out_str = liste_ID[0][i] + ' (ID : ' + str(liste_ID[1][i]) + ', réponse publique : ' + liste_ID[2][i] + ')'
            doc.add_paragraph(out_str, style='List Bullet')
        warning_fin(doc)
